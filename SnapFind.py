import os
import sys
import time
import pickle
import re
import tempfile
from pathlib import Path
from collections import defaultdict

FONT_SIZE_SMALL = 11
FONT_SIZE_NORMAL = 13
FONT_SIZE_LARGE = 16
FONT_SIZE_BOLD = 14

import cv2
import torch
import numpy as np
from PIL import Image, ImageSequence
from tqdm import tqdm
from transformers import SiglipModel, SiglipProcessor
from torch.utils.data import Dataset, DataLoader

def get_base_path():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))
BASE_PATH = get_base_path()

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QFileDialog, QCheckBox,
    QScrollArea, QFrame, QSplitter, QSizePolicy, QMessageBox,
    QGroupBox, QGridLayout, QTextEdit, QToolButton, QLabel, QDoubleSpinBox
)
from PyQt5.QtCore import (
    Qt, QThread, pyqtSignal, QTimer, QSize, QObject, pyqtSlot, QSettings
)
from PyQt5.QtGui import (
    QFont, QPixmap, QImage, QColor, QPalette, QInputMethodEvent,
    QIcon, QPainter, QBrush
)

if not hasattr(torch, 'get_default_device'):
    def get_default_device():
        return torch.device('cpu')
    torch.get_default_device = get_default_device

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx2txt import process as docx2txt_process
    DOCX2TXT_SUPPORT = True
except ImportError:
    DOCX2TXT_SUPPORT = False

class SafeLineEdit(QLineEdit):
    textChangedDelayed = pyqtSignal(str)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self._input_method_active = False
        self._text_change_timer = QTimer()
        self._text_change_timer.setSingleShot(True)
        self._text_change_timer.setInterval(200)
        self._text_change_timer.timeout.connect(self._emit_delayed_text_changed)
        self.textChanged.connect(self._on_text_changed)
        self.setFont(QFont("", FONT_SIZE_NORMAL))
    
    def inputMethodEvent(self, event: QInputMethodEvent):
        self._input_method_active = event.commitString() != "" or len(event.preeditString()) > 0
        super().inputMethodEvent(event)
    
    def _on_text_changed(self, text):
        if not self._input_method_active:
            self._text_change_timer.start()
    
    def _emit_delayed_text_changed(self):
        self.textChangedDelayed.emit(self.text())
    
    def get_real_text_length(self):
        return len(self.text())

class Config:
    SIGLIP_PATH = "./models"
    FINETUNED_MODEL_PATH = "./siglip2_pose_finetuned/checkpoint_epo0.pth"
    DEVICE = torch.device("cpu")
    TORCH_DTYPE = torch.float32
    
    IMAGE_EXTS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.jfif', '.gif'}
    VIDEO_EXTS = {'.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm'}
    TEXT_EXTS = {'.txt', '.md', '.csv', '.log', '.json', '.xml', '.py', '.cpp', '.h', '.java'}
    
    DOC_EXTS = {'.doc', '.docx'} if DOCX_SUPPORT else set()
    PDF_EXTS = {'.pdf'} if PDF_SUPPORT else set()
    ALL_DOC_EXTS = TEXT_EXTS.union(DOC_EXTS).union(PDF_EXTS)
    
    FEATURE_CACHE = "./multimodal_features.pkl"
    VIDEO_INTERVAL = 1
    TOP_K_PER_TYPE = 10
    BATCH_SIZE = 32
    CHECK_MODIFIED = True
    TEXT_MAX_LENGTH = 64
    
    ANCHOR_WINDOW_SIZE = 50
    CONTEXT_EXTENSION = 20
    ANCHOR_THRESHOLD = 0
    MAX_ANCHORS_PER_FILE = 20
    CACHE_VERSION = "1.4"
    
    VIDEO_DURATION_THRESHOLD = 45 * 60
    SHORT_VIDEO_INTERVAL = 1.0
    LONG_VIDEO_INTERVAL = 30.0

class DocumentParser:
    @staticmethod
    def extract_text_from_file(file_path, max_size_mb=10):
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        try:
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > max_size_mb:
                return ""
            
            if ext in Config.TEXT_EXTS:
                return DocumentParser._read_text_file(file_path)
            elif ext in Config.DOC_EXTS:
                return DocumentParser._read_doc_file(file_path)
            elif ext in Config.PDF_EXTS:
                return DocumentParser._read_pdf_file(file_path)
            else:
                return ""
        except Exception as e:
            return ""
    
    @staticmethod
    def _read_text_file(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            content = content.replace('\r\n', '\n').replace('\r', '\n')
            content = re.sub(r'\n{3,}', '\n\n', content)
            return content
        except Exception as e:
            return ""
    
    @staticmethod
    def _read_doc_file(file_path):
        if not DOCX_SUPPORT:
            return ""
        try:
            if DOCX2TXT_SUPPORT:
                try:
                    with tempfile.TemporaryDirectory() as tmpdir:
                        return docx2txt_process(str(file_path), tmpdir)
                except Exception:
                    pass
            doc = docx.Document(str(file_path))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return "\n".join(text_parts)
        except Exception:
            return ""
    
    @staticmethod
    def _read_pdf_file(file_path):
        if not PDF_SUPPORT:
            return ""
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                max_pages = min(50, len(pdf_reader.pages))
                for page_num in range(max_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text and text.strip():
                            text = text.replace('\r\n', '\n').replace('\r', '\n')
                            text = re.sub(r'\s+', ' ', text)
                            text = re.sub(r'\n{3,}', '\n\n', text)
                            text_parts.append(text.strip())
                    except Exception:
                        continue
            return "\n\n".join(text_parts)
        except Exception:
            return ""

def preprocess_text(text):
    if not text or not isinstance(text, str):
        return ""
    text = re.sub(r'[^\w\s\u4e00-\u9fff，。！？；：]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def segment_query(query):
    if not query:
        return []
    natural_segments = []
    if '的' in query:
        parts = query.split('的')
        for i in range(len(parts)-1):
            segment = parts[i] + '的' + parts[i+1]
            natural_segments.append(segment)
        natural_segments.extend(parts)
    patterns = [
        r'([一二三四五六七八九十两半几多全]+)([的]?)([\u4e00-\u9fff]+)',
        r'([\u4e00-\u9fff]+)([的]?)([\u4e00-\u9fff]+)',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, query)
        for match in matches:
            segment = ''.join(match)
            natural_segments.append(segment)
    segments = []
    query_len = len(query)
    for length in [query_len, 4, 3, 2]:
        if length <= query_len:
            for j in range(0, query_len - length + 1):
                segment = query[j:j+length]
                if '的' in segment and (segment[0] == '的' or segment[-1] == '的'):
                    continue
                if len(segment) >= 2:
                    segments.append(segment)
    segments.append(query)
    all_segments = natural_segments + segments
    all_segments = list(set(all_segments))
    
    def segment_score(seg):
        score = 0
        score += len(seg) * 100
        if '白色' in seg or '耳机' in seg or '有线' in seg:
            score += 50
        if '的' not in seg:
            score += 20
        if seg in query:
            score += 30
        return score
    
    all_segments.sort(key=lambda x: (-segment_score(x), x))
    return all_segments[:15]

class FeatureExtractor:
    def __init__(self, config):
        self.config = config
        self.processor = None
        self.model = None
        self.text_cache = {}
        self.doc_content_cache = {}
        self.initialized = False
        self.model_type = "Base Model"
    
    def initialize_model(self, status_callback=None):
        try:
            if status_callback:
                status_callback("Loading model...")
            
            self.processor = SiglipProcessor.from_pretrained(
                self.config.SIGLIP_PATH, 
                use_fast=True
            )
            self.model = SiglipModel.from_pretrained(
                self.config.SIGLIP_PATH,
                torch_dtype=self.config.TORCH_DTYPE,
                device_map=None,
                low_cpu_mem_usage=True
            ).to(self.config.DEVICE)
            
            self.model_type = "Base Model"
            if os.path.exists(self.config.FINETUNED_MODEL_PATH):
                try:
                    if status_callback:
                        status_callback("Fine-tuned model detected, loading...")
                    
                    checkpoint = torch.load(self.config.FINETUNED_MODEL_PATH, map_location='cpu')
                    model_state_dict = checkpoint['model_state_dict']
                    backbone_weights = {
                        k.replace('model.', ''): v for k, v in model_state_dict.items() 
                        if k.startswith('model.vision_model') or k.startswith('model.text_model')
                    }
                    self.model.load_state_dict(backbone_weights, strict=False)
                    self.model_type = "Fine-tuned Model"
                    if status_callback:
                        status_callback(f"✅ Successfully loaded fine-tuned weights: {self.config.FINETUNED_MODEL_PATH}")
                except Exception as e:
                    if status_callback:
                        status_callback(f"⚠️ Failed to load fine-tuned weights, using base model: {e}")
            else:
                if status_callback:
                    status_callback("Using base model")
            
            self.model.eval()
            self.initialized = True
            
            if status_callback:
                status_callback("✅ Model loaded successfully")
            
            return True
        except Exception as e:
            if status_callback:
                status_callback(f"❌ Model loading failed: {e}")
            return False
    
    def encode_text_batch(self, texts, batch_size=32):
        if not self.initialized:
            return np.array([])
        if isinstance(texts, str):
            texts = [texts]
        results = []
        to_encode = []
        cache_keys = []
        for text in texts:
            key = text[:100]
            if key in self.text_cache:
                results.append(self.text_cache[key])
            else:
                to_encode.append(text)
                cache_keys.append(key)
        if to_encode:
            all_features = []
            for i in range(0, len(to_encode), batch_size):
                batch_texts = to_encode[i:i+batch_size]
                inputs = self.processor(
                    text=batch_texts,
                    return_tensors="pt",
                    padding="max_length",
                    truncation=True,
                    max_length=self.config.TEXT_MAX_LENGTH
                )
                with torch.no_grad():
                    features = self.model.get_text_features(**inputs)
                    features = torch.nn.functional.normalize(features, p=2, dim=-1)
                    all_features.append(features.cpu())
            if all_features:
                encoded_features = torch.cat(all_features, dim=0).numpy()
                for idx, key in enumerate(cache_keys):
                    if idx < len(encoded_features):
                        self.text_cache[key] = encoded_features[idx]
                        results.append(encoded_features[idx])
        return np.array(results)
    
    def encode_text_single(self, text):
        if not self.initialized:
            return np.zeros(768)
        key = text[:100]
        if key in self.text_cache:
            return self.text_cache[key]
        features = self.encode_text_batch([text], batch_size=1)
        if len(features) > 0:
            self.text_cache[key] = features[0]
            return features[0]
        return np.zeros(768)
    
    def encode_images(self, image_list):
        if not self.initialized or not image_list:
            return np.array([])
        all_features = []
        batch_size = self.config.BATCH_SIZE
        for i in range(0, len(image_list), batch_size):
            batch_images = image_list[i:i + batch_size]
            pixel_values = []
            for img in batch_images:
                if isinstance(img, Image.Image) and getattr(img, "is_animated", False):
                    img.seek(0)
                inputs = self.processor(images=img, return_tensors="pt", padding=True)
                pixel_values.append(inputs["pixel_values"].squeeze(0))
            pixel_values = torch.stack(pixel_values).to(
                self.config.DEVICE, dtype=self.config.TORCH_DTYPE
            )
            with torch.no_grad():
                features = self.model.get_image_features(pixel_values=pixel_values)
                features = torch.nn.functional.normalize(features, p=2, dim=-1)
                all_features.append(features.cpu())
        return torch.cat(all_features, dim=0).numpy()

class VideoFrameDataset(Dataset):
    def __init__(self, frames, processor):
        self.frames = frames
        self.processor = processor

    def __len__(self):
        return len(self.frames)

    def __getitem__(self, idx):
        if idx < 0 or idx >= len(self.frames):
            return {"pixel_values": torch.zeros(3, 224, 224), "timestamp": 0.0}
        timestamp, frame = self.frames[idx]
        try:
            image = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            img_input = self.processor(
                images=image,
                return_tensors="pt",
                padding=True
            )
            return {
                "pixel_values": img_input["pixel_values"].squeeze(0),
                "timestamp": timestamp
            }
        except Exception as e:
            return {"pixel_values": torch.zeros(3, 224, 224), "timestamp": timestamp}

class MultimodalDataManager:
    def __init__(self, folder_path, config, extractor):
        self.folder_path = Path(folder_path)
        self.config = config
        self.extractor = extractor
        self.files = []
        self.existing_files = {}
        self.skipped_files = 0
        self.load_existing_features()
    
    def load_existing_features(self):
        if os.path.exists(self.config.FEATURE_CACHE):
            try:
                with open(self.config.FEATURE_CACHE, 'rb') as f:
                    data = pickle.load(f)
                if isinstance(data, dict) and 'files' in data:
                    self.files = data['files']
                else:
                    self.files = data
                
                valid_files = []
                for f in self.files:
                    if isinstance(f, dict) and 'feature' in f and 'path' in f:
                        if os.path.exists(f['path']):
                            feat = f['feature']
                            if isinstance(feat, np.ndarray) and len(feat.shape) > 0 and feat.shape[0] > 0:
                                valid_files.append(f)
                
                self.files = valid_files
                
                self.existing_files = {}
                for item in self.files:
                    path = item['path']
                    if item['type'] == 'image':
                        key = f"{path}_0"
                    elif item['type'] == 'video_frame':
                        frame_idx = item.get('metadata', {}).get('frame_idx', 0)
                        key = f"{path}_{frame_idx}"
                    self.existing_files[key] = item
            except Exception as e:
                self.files = []
                self.existing_files = {}
        else:
            self.files = []
            self.existing_files = {}
    
    def is_file_processed(self, file_path, metadata_key=""):
        unique_key = f"{str(file_path)}_{metadata_key}"
        
        if not os.path.exists(file_path):
            return False
        
        if unique_key not in self.existing_files:
            return False
        
        if self.config.CHECK_MODIFIED:
            try:
                existing = self.existing_files[unique_key]
                current_mtime = os.path.getmtime(file_path)
                stored_mtime = existing.get('metadata', {}).get('file_mtime', 0)
                if current_mtime > stored_mtime:
                    return False
            except Exception:
                return False
        
        return True
    
    def incremental_update(self, progress_callback=None, file_callback=None, search_types=None):
        if search_types is None:
            search_types = ['image', 'video_frame']
        
        self.cleanup_deleted_files()
        
        all_items = []
        video_frames = []
        self.skipped_files = 0
        
        try:
            file_list = list(self.folder_path.rglob("*"))
            total_files = len(file_list)
            
            if total_files == 0:
                if progress_callback:
                    progress_callback(100)
                return
            
            for idx, file_path in enumerate(file_list):
                if progress_callback and callable(progress_callback):
                    try:
                        progress_callback(int(idx / total_files * 50))
                    except Exception:
                        pass
                
                if not file_path.is_file():
                    continue
                
                ext = file_path.suffix.lower()
                
                process_image = 'image' in search_types and ext in self.config.IMAGE_EXTS
                process_video = 'video_frame' in search_types and ext in self.config.VIDEO_EXTS
                
                if not (process_image or process_video):
                    continue
                
                try:
                    file_mtime = os.path.getmtime(file_path)
                except Exception:
                    continue
                
                if process_image:
                    if not self.is_file_processed(file_path, 0):
                        if file_callback and callable(file_callback):
                            file_callback(f"Processing image: {file_path.name}")
                        
                        try:
                            img = Image.open(file_path).convert('RGB')
                            if ext == '.gif':
                                img.is_animated = True
                            all_items.append({
                                'type': 'image',
                                'data': img,
                                'metadata': {
                                    'path': str(file_path),
                                    'file_name': file_path.name,
                                    'file_mtime': file_mtime
                                }
                            })
                        except Exception as e:
                            self.skipped_files += 1
                            continue
                
                elif process_video:
                    try:
                        cap = cv2.VideoCapture(str(file_path))
                        if not cap.isOpened():
                            continue
                        
                        fps = cap.get(cv2.CAP_PROP_FPS)
                        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                        
                        if fps == 0 or total_frames == 0:
                            cap.release()
                            continue
                        
                        total_duration = total_frames / fps
                        
                        if total_duration > self.config.VIDEO_DURATION_THRESHOLD:
                            interval_seconds = self.config.LONG_VIDEO_INTERVAL
                        else:
                            interval_seconds = self.config.SHORT_VIDEO_INTERVAL
                        
                        frame_interval = max(1, int(fps * interval_seconds))
                        
                        need_processing = False
                        frames_to_process = []
                        
                        for current_frame in range(0, total_frames, frame_interval):
                            if not self.is_file_processed(file_path, current_frame):
                                need_processing = True
                                frames_to_process.append(current_frame)
                        
                        if need_processing:
                            if file_callback and callable(file_callback):
                                file_callback(f"Processing video: {file_path.name}")
                            
                            for current_frame in frames_to_process:
                                cap.set(cv2.CAP_PROP_POS_FRAMES, current_frame)
                                ret, frame = cap.read()
                                if not ret:
                                    continue
                                
                                timestamp = current_frame / fps
                                video_frames.append({
                                    'timestamp': timestamp,
                                    'frame': frame,
                                    'metadata': {
                                        'path': str(file_path),
                                        'timestamp': timestamp,
                                        'formatted_time': self._format_seconds(timestamp),
                                        'file_mtime': file_mtime,
                                        'frame_idx': current_frame,
                                        'interval_seconds': interval_seconds
                                    }
                                })
                        
                        cap.release()
                            
                    except Exception as e:
                        self.skipped_files += 1
                        continue
            
            if progress_callback and callable(progress_callback):
                progress_callback(50)
            
            if video_frames:
                try:
                    video_dataset = VideoFrameDataset(
                        [(item['timestamp'], item['frame']) for item in video_frames], 
                        self.extractor.processor
                    )
                    video_dataloader = DataLoader(
                        video_dataset, 
                        batch_size=self.config.BATCH_SIZE, 
                        num_workers=0
                    )
                    
                    video_features = []
                    total_batches = len(video_dataloader)
                    
                    with torch.no_grad():
                        for batch_idx, batch in enumerate(video_dataloader):
                            if progress_callback:
                                progress_callback(50 + int(batch_idx / total_batches * 25))
                            
                            pixel_values = batch["pixel_values"].to(
                                self.config.DEVICE, 
                                dtype=self.config.TORCH_DTYPE
                            )
                            features = self.extractor.model.get_image_features(pixel_values=pixel_values)
                            features = torch.nn.functional.normalize(features, p=2, dim=-1)
                            video_features.append(features.cpu().numpy())
                    
                    if video_features:
                        video_features = np.concatenate(video_features, axis=0)
                        for i, item in enumerate(video_frames):
                            if i < len(video_features):
                                all_items.append({
                                    'type': 'video_frame',
                                    'feature': video_features[i],
                                    'metadata': item['metadata']
                                })
                except Exception as e:
                    pass
            
            if progress_callback:
                progress_callback(75)
            
            image_items = [item for item in all_items if item['type'] == 'image']
            if image_items:
                try:
                    images = [item['data'] for item in image_items]
                    features = self.extractor.encode_images(images)
                    
                    for i, item in enumerate(image_items):
                        if i < len(features):
                            for j, a_item in enumerate(all_items):
                                if (a_item.get('type') == 'image' and 
                                    a_item.get('metadata', {}).get('path') == item['metadata']['path']):
                                    all_items[j]['feature'] = features[i]
                                    if 'data' in all_items[j]:
                                        del all_items[j]['data']
                                    break
                except Exception as e:
                    pass
            
            if progress_callback:
                progress_callback(90)
            
            new_features = []
            for item in all_items:
                if 'feature' in item and isinstance(item['feature'], np.ndarray):
                    new_features.append({
                        'type': item['type'],
                        'feature': item['feature'],
                        'path': item['metadata']['path'],
                        'metadata': item['metadata']
                    })
            
            updated_paths = set([item['path'] for item in new_features])
            self.files = [f for f in self.files if f['path'] not in updated_paths]
            self.files.extend(new_features)
            
            self.existing_files = {}
            for item in self.files:
                path = item['path']
                if item['type'] == 'image':
                    key = f"{path}_0"
                elif item['type'] == 'video_frame':
                    frame_idx = item.get('metadata', {}).get('frame_idx', 0)
                    key = f"{path}_{frame_idx}"
                self.existing_files[key] = item
            
            self.save_features()
            if progress_callback:
                progress_callback(100)
                
        except Exception as e:
            if progress_callback:
                progress_callback(0)
            raise
    
    def _format_seconds(self, seconds):
        total_seconds = int(round(seconds, 0))
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        secs = total_seconds % 60
        if hours > 0:
            return f"{hours}h{minutes:02d}m{secs:02d}s"
        else:
            return f"{minutes}m{secs:02d}s"
    
    def save_features(self):
        try:
            valid_files = []
            for f in self.files:
                if all(k in f for k in ['feature', 'path', 'type']):
                    feat = f['feature']
                    if isinstance(feat, np.ndarray) and len(feat.shape) == 1 and feat.shape[0] > 0:
                        valid_files.append(f)
            
            save_data = {
                'metadata': {
                    'version': self.config.CACHE_VERSION,
                    'created_at': time.time(),
                    'file_count': len(valid_files),
                    'image_count': len([f for f in valid_files if f['type'] == 'image']),
                    'video_frame_count': len([f for f in valid_files if f['type'] == 'video_frame'])
                },
                'files': valid_files
            }
            
            with open(self.config.FEATURE_CACHE, 'wb') as f:
                pickle.dump(save_data, f, protocol=pickle.HIGHEST_PROTOCOL)
        except Exception as e:
            raise
    
    def cleanup_deleted_files(self):
        original_count = len(self.files)
        self.files = [f for f in self.files if os.path.exists(f['path'])]
        deleted_count = original_count - len(self.files)
        if deleted_count > 0:
            self.save_features()
            self.existing_files = {}
            for item in self.files:
                path = item['path']
                if item['type'] == 'image':
                    key = f"{path}_0"
                elif item['type'] == 'video_frame':
                    frame_idx = item.get('metadata', {}).get('frame_idx', 0)
                    key = f"{path}_{frame_idx}"
                self.existing_files[key] = item

class FastSearchEngine:
    def __init__(self, config, extractor):
        self.config = config
        self.extractor = extractor
        self.files = []
        self.features = None
        self.doc_content_cache = {}
    
    def load_dataset(self, dataset):
        self.files = dataset.files
        if self.files:
            valid_features = []
            valid_files = []
            for f in self.files:
                if 'feature' in f:
                    feat = f['feature']
                    if not isinstance(feat, np.ndarray):
                        feat = np.array(feat)
                    if len(feat.shape) == 1 and feat.shape[0] > 0:
                        valid_features.append(feat)
                        valid_files.append(f)
            self.files = valid_files
            if valid_features:
                self.features = np.vstack(valid_features)
                self.features = self.features / (np.linalg.norm(self.features, axis=1, keepdims=True) + 1e-8)
    
    def search_media(self, query, search_types=None):
        if search_types is None:
            search_types = ['image', 'video_frame']
        if not self.files or self.features is None:
            return {'image': [], 'video_frame': []}
        query_feat = self.extractor.encode_text_single(preprocess_text(query))
        query_feat = query_feat / (np.linalg.norm(query_feat) + 1e-8)
        similarities = np.dot(self.features, query_feat.T).squeeze()
        for i, item in enumerate(self.files):
            item['similarity'] = float(similarities[i] if similarities.size > 1 else similarities)
        results = {}
        for file_type in ['image', 'video_frame']:
            if file_type in search_types:
                type_results = [f for f in self.files if f['type'] == file_type]
                type_results.sort(key=lambda x: x['similarity'], reverse=True)
                results[file_type] = type_results[:30]
            else:
                results[file_type] = []
        return results
    
    def smart_text_search(self, folder_path, query):
        query_len = len(query)
        query_segments = segment_query(query)
        if query_len <= 3:
            return self._char_level_search(folder_path, query)
        elif query_len <= 6:
            return self._hybrid_search(folder_path, query, query_segments)
        else:
            return self._semantic_anchor_search(folder_path, query, query_segments)
    
    def _char_level_search(self, folder_path, query):
        folder_path = Path(folder_path)
        results = []
        query_feat = self.extractor.encode_text_single(query)
        query_feat = query_feat / (np.linalg.norm(query_feat) + 1e-8)
        for file_path in folder_path.rglob("*"):
            ext = file_path.suffix.lower()
            if ext not in self.config.ALL_DOC_EXTS:
                continue
            try:
                content = self._get_document_content(file_path)
                if not content:
                    continue
                for char in query:
                    start = 0
                    char_count = 0
                    while char_count < 5:
                        pos = content.find(char, start)
                        if pos == -1:
                            break
                        window_start = max(0, pos - 40)
                        window_end = min(len(content), pos + 40)
                        window_text = content[window_start:window_end]
                        window_feat = self.extractor.encode_text_single(window_text)
                        window_feat = window_feat / (np.linalg.norm(window_feat) + 1e-8)
                        similarity = np.dot(query_feat, window_feat.T)
                        line_number = content[:pos].count('\n') + 1
                        results.append({
                            'type': 'text_sentence',
                            'file_type': ext,
                            'similarity': float(similarity),
                            'path': str(file_path),
                            'line_number': line_number,
                            'sentence_content': window_text[:150]
                        })
                        start = pos + 1
                        char_count += 1
            except Exception:
                continue
        return self._get_top_k_results(results, k=30)
    
    def _hybrid_search(self, folder_path, query, query_segments):
        folder_path = Path(folder_path)
        results = []
        all_texts = [query] + query_segments
        all_features = self.extractor.encode_text_batch(all_texts)
        if len(all_features) == 0:
            return []
        query_feat = all_features[0]
        query_feat = query_feat / (np.linalg.norm(query_feat) + 1e-8)
        for file_path in folder_path.rglob("*"):
            ext = file_path.suffix.lower()
            if ext not in self.config.ALL_DOC_EXTS:
                continue
            try:
                content = self._get_document_content(file_path)
                if not content:
                    continue
                for segment in query_segments[:5]:
                    if len(segment) < 2:
                        continue
                    start = 0
                    seg_count = 0
                    while seg_count < 3:
                        pos = content.find(segment, start)
                        if pos == -1:
                            break
                        context_start = max(0, pos - 50)
                        context_end = min(len(content), pos + len(segment) + 50)
                        context_text = content[context_start:context_end]
                        context_feat = self.extractor.encode_text_single(context_text)
                        context_feat = context_feat / (np.linalg.norm(context_feat) + 1e-8)
                        similarity = np.dot(query_feat, context_feat.T)
                        line_number = content[:pos].count('\n') + 1
                        results.append({
                            'type': 'text_sentence',
                            'file_type': ext,
                            'similarity': float(similarity),
                            'path': str(file_path),
                            'line_number': line_number,
                            'sentence_content': context_text[:150]
                        })
                        start = pos + 1
                        seg_count += 1
            except Exception:
                continue
        return self._get_top_k_results(results, k=30)
    
    def _semantic_anchor_search(self, folder_path, query, query_segments):
        folder_path = Path(folder_path)
        results = []
        all_texts = [query] + query_segments
        all_features = self.extractor.encode_text_batch(all_texts)
        if len(all_features) == 0:
            return []
        query_feat = all_features[0]
        query_feat = query_feat / (np.linalg.norm(query_feat) + 1e-8)
        segment_features = all_features[1:]
        for i in range(len(segment_features)):
            segment_features[i] = segment_features[i] / (np.linalg.norm(segment_features[i]) + 1e-8)
        for file_path in folder_path.rglob("*"):
            ext = file_path.suffix.lower()
            if ext not in self.config.ALL_DOC_EXTS:
                continue
            try:
                content = self._get_document_content(file_path)
                if not content:
                    continue
                anchor_candidates = []
                for segment_idx, segment in enumerate(query_segments):
                    if len(segment) < 2:
                        continue
                    start = 0
                    segment_positions = []
                    while True:
                        pos = content.find(segment, start)
                        if pos == -1:
                            break
                        segment_positions.append(pos)
                        start = pos + 1
                    segment_positions = segment_positions[:5]
                    for pos in segment_positions:
                        window_start = max(0, pos - self.config.ANCHOR_WINDOW_SIZE)
                        window_end = min(len(content), pos + len(segment) + self.config.ANCHOR_WINDOW_SIZE)
                        anchor_text = content[window_start:window_end]
                        anchor_candidates.append({
                            'position': pos,
                            'segment_idx': segment_idx,
                            'anchor_text': anchor_text,
                            'window_start': window_start,
                            'window_end': window_end
                        })
                if not anchor_candidates:
                    continue
                anchor_candidates = anchor_candidates[:self.config.MAX_ANCHORS_PER_FILE * 3]
                anchor_texts = [candidate['anchor_text'] for candidate in anchor_candidates]
                anchor_features = self.extractor.encode_text_batch(anchor_texts)
                if len(anchor_features) == 0:
                    continue
                for i in range(len(anchor_features)):
                    anchor_features[i] = anchor_features[i] / (np.linalg.norm(anchor_features[i]) + 1e-8)
                valid_anchors = []
                for idx, candidate in enumerate(anchor_candidates):
                    if idx >= len(anchor_features):
                        continue
                    segment_idx = candidate['segment_idx']
                    if segment_idx >= len(segment_features):
                        continue
                    anchor_sim = np.dot(anchor_features[idx], segment_features[segment_idx].T)
                    if anchor_sim > self.config.ANCHOR_THRESHOLD:
                        valid_anchors.append({
                            'position': candidate['position'],
                            'window_start': candidate['window_start'],
                            'window_end': candidate['window_end']
                        })
                if not valid_anchors:
                    continue
                for anchor in valid_anchors[:self.config.MAX_ANCHORS_PER_FILE]:
                    context_start = max(0, anchor['window_start'] - self.config.CONTEXT_EXTENSION)
                    context_end = min(len(content), anchor['window_end'] + self.config.CONTEXT_EXTENSION)
                    context_text = content[context_start:context_end]
                    context_feat = self.extractor.encode_text_single(context_text)
                    context_feat = context_feat / (np.linalg.norm(context_feat) + 1e-8)
                    final_sim = np.dot(query_feat, context_feat.T)
                    line_number = content[:anchor['position']].count('\n') + 1
                    results.append({
                        'type': 'text_sentence',
                        'file_type': ext,
                        'similarity': float(final_sim),
                        'path': str(file_path),
                        'line_number': line_number,
                        'sentence_content': context_text[:200]
                    })
            except Exception:
                continue
        return self._get_top_k_results(results, k=30)
    
    def _get_document_content(self, file_path):
        file_path_str = str(file_path)
        if file_path_str in self.doc_content_cache:
            return self.doc_content_cache[file_path_str]
        content = DocumentParser.extract_text_from_file(file_path)
        self.doc_content_cache[file_path_str] = content
        return content
    
    def _get_top_k_results(self, results, k=30):
        if not results:
            return []
        deduped = []
        seen_keys = set()
        for result in results:
            key = f"{result['path']}:{result['line_number']}:{result['sentence_content'][:50]}"
            if key not in seen_keys:
                seen_keys.add(key)
                deduped.append(result)
        deduped.sort(key=lambda x: x['similarity'], reverse=True)
        return deduped[:k]

class SearchThread(QThread):
    search_progress = pyqtSignal(int)
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)
    current_file = pyqtSignal(str)
    search_status = pyqtSignal(str)
    
    def __init__(self, folder_path, query, search_types, extractor, config):
        super().__init__()
        self.folder_path = folder_path
        self.query = query
        self.search_types = search_types
        self.extractor = extractor
        self.config = config
    
    def run(self):
        try:
            self.search_status.emit("Search started...")
            self.search_progress.emit(0)
            
            dataset = MultimodalDataManager(self.folder_path, self.config, self.extractor)
            
            media_types = [t for t in self.search_types if t in ['image', 'video_frame']]
            if media_types:
                self.search_progress.emit(5)
                dataset.incremental_update(
                    progress_callback=lambda x: self._handle_progress(x, "Search progress"),
                    file_callback=lambda x: self._handle_file(x),
                    search_types=media_types
                )
            
            engine = FastSearchEngine(self.config, self.extractor)
            if media_types:
                engine.load_dataset(dataset)
            
            results = {'image': [], 'video_frame': [], 'text_sentence': []}
            self.search_progress.emit(70)
            
            if media_types:
                self.search_status.emit("Comparing features...")
                media_results = engine.search_media(self.query, media_types)
                results['image'] = media_results.get('image', [])
                results['video_frame'] = media_results.get('video_frame', [])
            
            if 'text_sentence' in self.search_types:
                self.search_status.emit("Searching documents...")
                text_results = engine.smart_text_search(self.folder_path, self.query)
                results['text_sentence'] = text_results
            
            self.search_progress.emit(100)
            self.search_status.emit("✅ Search completed")
            self.finished.emit(results)
            
        except Exception as e:
            import traceback
            error_detail = f"{str(e)}\n{traceback.format_exc()}"
            print(f"Search thread error: {error_detail}")
            self.search_status.emit("❌ Search error")
            self.error.emit(str(e))
    
    def _handle_progress(self, value, prefix):
        if hasattr(self, 'current_processing_file') and self.current_processing_file:
            self.search_status.emit(f"{prefix}... {value}% ({self.current_processing_file})")
        else:
            self.search_status.emit(f"{prefix}... {value}%")
        self.search_progress.emit(value)
    
    def _handle_file(self, file_name):
        import os
        if len(file_name) > 30:
            file_name = "..." + file_name[-30:]
        self.current_processing_file = file_name
        self.current_file.emit(file_name)

class ResultItemWidget(QWidget):
    def __init__(self, result, parent=None):
        super().__init__(parent)
        self.result = result
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        self.setLayout(layout)
        self.setStyleSheet("""
            QWidget {
                border: 1px solid #cccccc;
                border-radius: 5px;
                padding: 8px;
                margin: 5px;
                background-color: white;
            }
            QWidget:hover {
                background-color: #f5f5f5;
            }
        """)
        
        similarity = self.result.get('similarity', 0.0)
        similarity_label = QLabel(f"Similarity: {similarity:.4f}")
        similarity_label.setFont(QFont("", FONT_SIZE_BOLD, QFont.Bold))
        similarity_label.setStyleSheet("font-weight: bold;")
        layout.addWidget(similarity_label)
        
        path_label = QLabel(f"{self.result.get('path', '')}")
        path_label.setFont(QFont("", FONT_SIZE_NORMAL))
        path_label.setWordWrap(True)
        path_label.setStyleSheet("color: #666666;")
        path_label.setTextInteractionFlags(Qt.TextSelectableByMouse | Qt.TextSelectableByKeyboard)
        path_label.setCursor(Qt.IBeamCursor)
        layout.addWidget(path_label)
        
        result_type = self.result.get('type', '')
        if result_type == 'image':
            self.add_image_content(layout)
        elif result_type == 'video_frame':
            self.add_video_content(layout)
        elif result_type == 'text_sentence':
            self.add_text_content(layout)
    
    def add_image_content(self, layout):
        try:
            img_path = self.result.get('path', '')
            pixmap = QPixmap(img_path)
            if not pixmap.isNull():
                pixmap = pixmap.scaled(200, 150, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                img_label = QLabel()
                img_label.setPixmap(pixmap)
                img_label.setAlignment(Qt.AlignCenter)
                layout.addWidget(img_label)
        except Exception:
            pass
    
    def add_video_content(self, layout):
        time_str = self.result.get('metadata', {}).get('formatted_time', '')
        time_label = QLabel(f"Timestamp: {time_str}")
        time_label.setFont(QFont("", FONT_SIZE_NORMAL))
        layout.addWidget(time_label)
        try:
            cap = cv2.VideoCapture(self.result.get('path', ''))
            frame_idx = self.result.get('metadata', {}).get('frame_idx', 0)
            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
            ret, frame = cap.read()
            if ret:
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                h, w, ch = frame.shape
                bytes_per_line = ch * w
                qt_image = QImage(frame.data, w, h, bytes_per_line, QImage.Format_RGB888)
                pixmap = QPixmap.fromImage(qt_image)
                pixmap = pixmap.scaled(200, 150, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                img_label = QLabel()
                img_label.setPixmap(pixmap)
                img_label.setAlignment(Qt.AlignCenter)
                layout.addWidget(img_label)
            cap.release()
        except Exception:
            pass
    
    def add_text_content(self, layout):
        line_num = self.result.get('line_number', 0)
        line_label = QLabel(f"Line {line_num}")
        line_label.setFont(QFont("", FONT_SIZE_NORMAL))
        layout.addWidget(line_label)
        content = self.result.get('sentence_content', '')
        text_edit = QTextEdit()
        text_edit.setText(content)
        text_edit.setFont(QFont("", FONT_SIZE_NORMAL))
        text_edit.setReadOnly(True)
        text_edit.setMaximumHeight(100)
        layout.addWidget(text_edit)

class ResultPageWidget(QWidget):
    def __init__(self, result_type, parent=None, main_window=None):
        super().__init__(parent)
        self.result_type = result_type
        self.main_window = main_window
        self.current_page = 0
        self.results_per_page = 10
        self.all_results = []
        self.is_maximized = False
        self.init_ui()
    
    def init_ui(self):
        self.main_layout = QVBoxLayout()
        self.setLayout(self.main_layout)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        
        title_bar = QWidget()
        title_layout = QHBoxLayout(title_bar)
        title_layout.setContentsMargins(5, 5, 5, 5)
        title_layout.setSpacing(5)
        
        type_names = {
            'image': '🖼️ Image Results',
            'video_frame': '🎬 Video Results',
            'text_sentence': '📄 Document Results'
        }
        self.title_label = QLabel(type_names.get(self.result_type, 'Results'))
        self.title_label.setFont(QFont("", FONT_SIZE_LARGE, QFont.Bold))
        self.title_label.setStyleSheet("padding: 5px;")
        title_layout.addWidget(self.title_label)
        title_layout.addStretch()
        
        self.max_btn = QPushButton("Maximize")
        self.max_btn.setToolTip("Maximize current panel")
        self.max_btn.clicked.connect(self.toggle_maximize)
        self.max_btn.setFont(QFont("", FONT_SIZE_NORMAL))
        self.max_btn.setFixedSize(80, 30)
        self.max_btn.setStyleSheet("""
            QPushButton {
                border: 1px solid #cccccc;
                border-radius: 3px;
                background-color: #f8f8f8;
            }
            QPushButton:hover {
                background-color: #e8e8e8;
            }
        """)
        title_layout.addWidget(self.max_btn)
        
        self.close_btn = QPushButton("Close")
        self.close_btn.setToolTip("Restore original layout")
        self.close_btn.clicked.connect(self.restore_layout)
        self.close_btn.hide()
        self.close_btn.setFont(QFont("", FONT_SIZE_NORMAL))
        self.close_btn.setFixedSize(80, 30)
        self.close_btn.setStyleSheet("""
            QPushButton {
                border: 1px solid #cccccc;
                border-radius: 3px;
                background-color: #f8f8f8;
            }
            QPushButton:hover {
                background-color: #e8e8e8;
            }
        """)
        title_layout.addWidget(self.close_btn)
        
        self.main_layout.addWidget(title_bar)
        
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_content = QWidget()
        self.results_layout = QGridLayout(self.scroll_content)
        self.scroll_area.setWidget(self.scroll_content)
        self.main_layout.addWidget(self.scroll_area, stretch=1)
        
        self.pagination_layout = QHBoxLayout()
        self.prev_btn = QPushButton("Previous")
        self.prev_btn.setFont(QFont("", FONT_SIZE_NORMAL))
        self.prev_btn.clicked.connect(self.prev_page)
        self.prev_btn.setEnabled(False)
        self.pagination_layout.addWidget(self.prev_btn)
        
        self.page_label = QLabel("Page 0 / 0")
        self.page_label.setFont(QFont("", FONT_SIZE_NORMAL))
        self.page_label.setAlignment(Qt.AlignCenter)
        self.pagination_layout.addWidget(self.page_label)
        
        self.next_btn = QPushButton("Next")
        self.next_btn.setFont(QFont("", FONT_SIZE_NORMAL))
        self.next_btn.clicked.connect(self.next_page)
        self.next_btn.setEnabled(False)
        self.pagination_layout.addWidget(self.next_btn)
        
        self.main_layout.addLayout(self.pagination_layout)
    
    def toggle_maximize(self):
        if not self.is_maximized:
            self.maximize()
        else:
            self.restore_layout()
    
    def maximize(self):
        if not self.main_window:
            return
        self.is_maximized = True
        
        if self.result_type == 'video_frame':
            self.main_window.image_results.hide()
            self.main_window.text_results.hide()
        elif self.result_type == 'image':
            self.main_window.video_results.hide()
            self.main_window.text_results.hide()
        elif self.result_type == 'text_sentence':
            self.main_window.video_results.hide()
            self.main_window.image_results.hide()
        
        self.max_btn.hide()
        self.close_btn.show()
        
        self.main_window.splitter.setSizes([100, 0, 0])
        
        self.update_display()
    
    def restore_layout(self):
        if not self.main_window:
            return
        self.is_maximized = False
        
        if self.main_window.video_check.isChecked():
            self.main_window.video_results.show()
        if self.main_window.image_check.isChecked():
            self.main_window.image_results.show()
        if self.main_window.text_check.isChecked():
            self.main_window.text_results.show()
        
        self.max_btn.show()
        self.close_btn.hide()
        
        selected_types = []
        if self.main_window.video_check.isChecked():
            selected_types.append('video_frame')
        if self.main_window.image_check.isChecked():
            selected_types.append('image')
        if self.main_window.text_check.isChecked():
            selected_types.append('text_sentence')
        
        if len(selected_types) == 1:
            self.main_window.splitter.setSizes([100, 0, 0])
        elif len(selected_types) == 2:
            self.main_window.splitter.setSizes([50, 50, 0])
        else:
            self.main_window.splitter.setSizes([33, 33, 34])
        
        self.update_display()
    
    def set_results(self, results):
        self.all_results = results
        self.current_page = 0
        self.update_display()
    
    def update_display(self):
        for i in reversed(range(self.results_layout.count())):
            widget = self.results_layout.itemAt(i).widget()
            if widget:
                widget.deleteLater()
        
        total_pages = max(1, (len(self.all_results) + self.results_per_page - 1) // self.results_per_page)
        self.page_label.setText(f"Page {self.current_page + 1} / {total_pages}")
        self.prev_btn.setEnabled(self.current_page > 0)
        self.next_btn.setEnabled(self.current_page < total_pages - 1)
        
        start_idx = self.current_page * self.results_per_page
        end_idx = min(start_idx + self.results_per_page, len(self.all_results))
        current_results = self.all_results[start_idx:end_idx]
        
        cols = 4 if self.is_maximized else 3
        for i, result in enumerate(current_results):
            row = i // cols
            col = i % cols
            item_widget = ResultItemWidget(result)
            self.results_layout.addWidget(item_widget, row, col)
    
    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.update_display()
    
    def next_page(self):
        total_pages = (len(self.all_results) + self.results_per_page - 1) // self.results_per_page
        if self.current_page < total_pages - 1:
            self.current_page += 1
            self.update_display()

class MultimodalSearchUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.config = Config()
        self.extractor = FeatureExtractor(self.config)
        self.search_thread = None
        self.settings = QSettings("MultimodalSearch", "SearchApp")
        
        self.current_processing_file = ""
        self.current_progress = 0
        self.is_searching = False
        
        self.init_ui()
        self.init_model()
        self.load_last_settings()

    def init_ui(self):
        self.setWindowTitle("SnapFind v1.4 (CPU Version)")
        self.setMinimumSize(1300, 800)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)

        control_panel = QWidget()
        control_layout = QVBoxLayout(control_panel)
        control_panel.setStyleSheet("background-color: #f0f0f0; padding: 10px; border-radius: 5px;")
        control_layout.setContentsMargins(10, 10, 10, 10)
        control_layout.setSpacing(8)

        query_layout = QHBoxLayout()
        query_layout.setSpacing(8)
        
        query_label = QLabel("Search Keywords:")
        query_label.setFont(QFont("", FONT_SIZE_NORMAL, QFont.Bold))
        query_label.setMinimumWidth(80)
        query_layout.addWidget(query_label)
        
        self.query_edit = SafeLineEdit()
        self.query_edit.setPlaceholderText(f"Enter search keywords (max {self.config.TEXT_MAX_LENGTH} characters)")
        self.query_edit.setMinimumHeight(35)
        self.query_edit.textChangedDelayed.connect(self.check_input_length)
        query_layout.addWidget(self.query_edit, stretch=1)
        
        self.length_warning = QLabel("")
        self.length_warning.setFont(QFont("", FONT_SIZE_SMALL))
        self.length_warning.setStyleSheet("color: red;")
        self.length_warning.setMinimumWidth(120)
        query_layout.addWidget(self.length_warning)
        control_layout.addLayout(query_layout)

        path_interval_layout = QHBoxLayout()
        path_interval_layout.setSpacing(15)

        folder_container = QWidget()
        folder_layout = QHBoxLayout(folder_container)
        folder_layout.setContentsMargins(0, 0, 0, 0)
        folder_layout.setSpacing(8)
        
        folder_label = QLabel("Search Path:")
        folder_label.setFont(QFont("", FONT_SIZE_NORMAL, QFont.Bold))
        folder_label.setMinimumWidth(80)
        folder_layout.addWidget(folder_label)
        
        self.folder_edit = QLineEdit()
        self.folder_edit.setFont(QFont("", FONT_SIZE_NORMAL))
        self.folder_edit.setPlaceholderText("Select folder to search...")
        self.folder_edit.setMinimumHeight(35)
        self.folder_edit.setMinimumWidth(300)
        folder_layout.addWidget(self.folder_edit, stretch=1)
        
        folder_btn = QPushButton("Browse")
        folder_btn.setFont(QFont("", FONT_SIZE_NORMAL))
        folder_btn.setFixedSize(70, 35)
        folder_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 3px;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        folder_btn.clicked.connect(self.select_folder)
        folder_layout.addWidget(folder_btn)
        
        path_interval_layout.addWidget(folder_container, stretch=1)

        separator = QFrame()
        separator.setFrameShape(QFrame.VLine)
        separator.setFrameShadow(QFrame.Sunken)
        separator.setFixedWidth(1)
        separator.setStyleSheet("background-color: #cccccc;")
        path_interval_layout.addWidget(separator)

        interval_container = QWidget()
        interval_layout = QHBoxLayout(interval_container)
        interval_layout.setContentsMargins(5, 0, 5, 0)
        interval_layout.setSpacing(8)
        
        interval_label = QLabel("Video Frame Interval:")
        interval_label.setFont(QFont("", FONT_SIZE_NORMAL, QFont.Bold))
        interval_layout.addWidget(interval_label)
        
        short_label = QLabel("≤45 min:")
        short_label.setFont(QFont("", FONT_SIZE_SMALL))
        interval_layout.addWidget(short_label)
        
        self.short_interval_spin = QDoubleSpinBox()
        self.short_interval_spin.setFont(QFont("", FONT_SIZE_NORMAL))
        self.short_interval_spin.setRange(0.5, 60.0)
        self.short_interval_spin.setSingleStep(0.5)
        self.short_interval_spin.setValue(self.config.SHORT_VIDEO_INTERVAL)
        self.short_interval_spin.setSuffix("s")
        self.short_interval_spin.setToolTip("Frame interval for short videos (≤45 min)")
        interval_layout.addWidget(self.short_interval_spin)
        
        long_label = QLabel(">45 min:")
        long_label.setFont(QFont("", FONT_SIZE_SMALL))
        interval_layout.addWidget(long_label)
        
        self.long_interval_spin = QDoubleSpinBox()
        self.long_interval_spin.setFont(QFont("", FONT_SIZE_NORMAL))
        self.long_interval_spin.setRange(1.0, 300.0)
        self.long_interval_spin.setSingleStep(1.0)
        self.long_interval_spin.setValue(self.config.LONG_VIDEO_INTERVAL)
        self.long_interval_spin.setSuffix("s")
        self.long_interval_spin.setToolTip("Frame interval for long videos (>45 min)")
        interval_layout.addWidget(self.long_interval_spin)
        
        interval_layout.addStretch()
        path_interval_layout.addWidget(interval_container, stretch=1)
        
        control_layout.addLayout(path_interval_layout)

        type_button_layout = QHBoxLayout()
        type_button_layout.setSpacing(15)
        
        type_container = QWidget()
        type_layout = QHBoxLayout(type_container)
        type_layout.setContentsMargins(5, 0, 5, 0)
        type_layout.setSpacing(15)
        
        type_label = QLabel("Search Type:")
        type_label.setFont(QFont("", FONT_SIZE_NORMAL, QFont.Bold))
        type_label.setMinimumWidth(75)
        type_layout.addWidget(type_label)
        
        self.video_check = QCheckBox("Video")
        self.video_check.setFont(QFont("", FONT_SIZE_NORMAL))
        self.video_check.setChecked(True)
        
        self.image_check = QCheckBox("Images")
        self.image_check.setFont(QFont("", FONT_SIZE_NORMAL))
        
        self.text_check = QCheckBox("Documents")
        self.text_check.setFont(QFont("", FONT_SIZE_NORMAL))
        
        type_layout.addWidget(self.video_check)
        type_layout.addWidget(self.image_check)
        type_layout.addWidget(self.text_check)
        type_layout.addStretch()
        
        self.search_btn = QPushButton("Start")
        self.search_btn.setFont(QFont("", FONT_SIZE_NORMAL, QFont.Bold))
        self.search_btn.setFixedSize(90, 35)
        self.search_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #0b7dda;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
        """)
        self.search_btn.clicked.connect(self.start_search)
        type_layout.addWidget(self.search_btn)
        
        type_button_layout.addWidget(type_container, stretch=1)
        control_layout.addLayout(type_button_layout)
        
        main_layout.addWidget(control_panel)

        self.splitter = QSplitter(Qt.Horizontal)
        self.splitter.setStyleSheet("QSplitter::handle { background-color: #cccccc; }")
        main_layout.addWidget(self.splitter, stretch=1)
        
        self.video_results = ResultPageWidget('video_frame', main_window=self)
        self.image_results = ResultPageWidget('image', main_window=self)
        self.text_results = ResultPageWidget('text_sentence', main_window=self)
        
        self.video_results.hide()
        self.image_results.hide()
        self.text_results.hide()
        
        self.splitter.addWidget(self.video_results)
        self.splitter.addWidget(self.image_results)
        self.splitter.addWidget(self.text_results)

        status_container = QWidget()
        status_container.setStyleSheet("background-color: #f0f0f0; border-top: 1px solid #cccccc;")
        status_layout = QHBoxLayout(status_container)
        status_layout.setContentsMargins(10, 5, 10, 5)
        status_layout.setSpacing(5)
        
        self.status_label = QLabel("Ready")
        self.status_label.setFont(QFont("", FONT_SIZE_SMALL))
        self.status_label.setStyleSheet("color: #333333; padding: 2px 5px;")
        status_layout.addWidget(self.status_label)
        
        self.progress_bar = QLabel("")
        self.progress_bar.setFont(QFont("", FONT_SIZE_SMALL))
        self.progress_bar.setStyleSheet("color: #666666; padding: 2px 5px;")
        status_layout.addWidget(self.progress_bar)
        
        status_layout.addStretch()
        
        main_layout.addWidget(status_container)

    def init_model(self):
        QTimer.singleShot(100, lambda: self.extractor.initialize_model(
            status_callback=lambda msg: self.status_label.setText(msg)
        ))

    def select_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, "Select Folder")
        if folder_path:
            self.folder_edit.setText(folder_path)

    def check_input_length(self):
        length = self.query_edit.get_real_text_length()
        if length > self.config.TEXT_MAX_LENGTH:
            text = self.query_edit.text()
            self.query_edit.setText(text[:self.config.TEXT_MAX_LENGTH])
            self.length_warning.setText(f"⚠️ Input cannot exceed {self.config.TEXT_MAX_LENGTH} characters!")
        else:
            self.length_warning.setText(f"Characters: {length}/{self.config.TEXT_MAX_LENGTH}")
            if length == 0:
                self.length_warning.setText("")

    def start_search(self):
        folder_path = self.folder_edit.text().strip()
        query = self.query_edit.text().strip()
        
        self.save_current_settings()
        
        if not folder_path:
            QMessageBox.warning(self, "Warning", "Please select a search folder first!")
            return
        if not query:
            QMessageBox.warning(self, "Warning", "Please enter search keywords!")
            return
        
        search_types = []
        if self.video_check.isChecked():
            search_types.append('video_frame')
        if self.image_check.isChecked():
            search_types.append('image')
        if self.text_check.isChecked():
            search_types.append('text_sentence')
        
        if not search_types:
            QMessageBox.warning(self, "Warning", "Please select at least one search type!")
            return
        
        if self.search_thread and self.search_thread.isRunning():
            self.search_thread.terminate()
            self.search_thread.wait()
        
        self.current_processing_file = ""
        self.current_progress = 0
        self.is_searching = True
        
        self.status_label.setText("Search started...")
        self.progress_bar.setText("")
        
        self.config.SHORT_VIDEO_INTERVAL = self.short_interval_spin.value()
        self.config.LONG_VIDEO_INTERVAL = self.long_interval_spin.value()
        
        self.search_thread = SearchThread(folder_path, query, search_types, self.extractor, self.config)
        
        self.search_thread.search_progress.connect(self.update_progress)
        self.search_thread.finished.connect(self.on_search_finished)
        self.search_thread.error.connect(self.on_search_error)
        self.search_thread.current_file.connect(self.update_current_file)
        self.search_thread.search_status.connect(self.update_status_label)
        
        self.search_btn.setEnabled(False)
        self.search_thread.start()

    def update_progress(self, value):
        self.current_progress = value
        self.progress_bar.setText(f"Progress: {value}%")

    def update_current_file(self, file_name):
        if len(file_name) > 30:
            file_name = "..." + file_name[-30:]
        self.current_processing_file = file_name
        self.status_label.setText(f"Processing: {file_name}")

    def update_status_label(self, message):
        self.status_label.setText(message)

    def on_search_finished(self, results):
        self.search_btn.setEnabled(True)
        self.is_searching = False
        self.current_processing_file = ""
        self.current_progress = 0
        self.progress_bar.setText("")
        
        selected_types = []
        if self.video_check.isChecked():
            selected_types.append('video_frame')
        if self.image_check.isChecked():
            selected_types.append('image')
        if self.text_check.isChecked():
            selected_types.append('text_sentence')
        
        self.video_results.setVisible('video_frame' in selected_types)
        self.image_results.setVisible('image' in selected_types)
        self.text_results.setVisible('text_sentence' in selected_types)
        
        self.video_results.set_results(results['video_frame'])
        self.image_results.set_results(results['image'])
        self.text_results.set_results(results['text_sentence'])
        
        if len(selected_types) == 1:
            self.splitter.setSizes([100, 0, 0])
        elif len(selected_types) == 2:
            self.splitter.setSizes([50, 50, 0])
        else:
            self.splitter.setSizes([33, 33, 34])
        
        self.status_label.setText("✅ Search completed")

    def on_search_error(self, error_msg):
        self.search_btn.setEnabled(True)
        self.is_searching = False
        self.current_processing_file = ""
        self.current_progress = 0
        self.progress_bar.setText("")
        self.status_label.setText(f"❌ Search error: {error_msg[:50]}...")
        QMessageBox.critical(self, "Error", f"An error occurred during search: {error_msg}")

    def load_last_settings(self):
        last_folder = self.settings.value("last_folder", "")
        last_query = self.settings.value("last_query", "")
        last_short_interval = self.settings.value("short_interval", 1.0, type=float)
        last_long_interval = self.settings.value("long_interval", 30.0, type=float)
        
        if last_folder:
            self.folder_edit.setText(last_folder)
        if last_query:
            self.query_edit.setText(last_query)
            self.check_input_length()
        
        self.short_interval_spin.setValue(last_short_interval)
        self.long_interval_spin.setValue(last_long_interval)

    def save_current_settings(self):
        current_folder = self.folder_edit.text().strip()
        current_query = self.query_edit.text().strip()
        current_short_interval = self.short_interval_spin.value()
        current_long_interval = self.long_interval_spin.value()
        
        self.settings.setValue("last_folder", current_folder)
        self.settings.setValue("last_query", current_query)
        self.settings.setValue("short_interval", current_short_interval)
        self.settings.setValue("long_interval", current_long_interval)

if __name__ == "__main__":
    import sys
    from PyQt5.QtCore import Qt
    
    if hasattr(Qt, 'AA_EnableHighDpiScaling'):
        QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
    if hasattr(Qt, 'AA_UseHighDpiPixmaps'):
        QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
    
    app = QApplication(sys.argv)
    
    font = QFont()
    font.setFamily("SimHei" if sys.platform == "win32" else "WenQuanYi Micro Hei")
    app.setFont(font)
    
    window = MultimodalSearchUI()
    window.show()
    
    sys.exit(app.exec_())